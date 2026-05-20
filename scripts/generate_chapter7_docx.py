"""将第7章测试论文markdown转换为docx格式。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
INPUT_PATH = PROJECT_ROOT / "论文_第7章_系统测试.md"
OUTPUT_PATH = PROJECT_ROOT / "论文_第7章_系统测试.docx"


def set_cell_font(cell, font_name="宋体", font_size=Pt(10.5)):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(doc, text, first_line_indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return para


def add_table_caption(doc, caption_text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(caption_text)
    run.font.name = "黑体"
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_font(cell, font_name="黑体", font_size=Pt(10.5))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = value
            set_cell_font(cell, font_name="宋体", font_size=Pt(10.5))

    doc.add_paragraph()


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def parse_md_table(lines):
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        row = [c.strip() for c in line.strip("|").split("|")]
        rows.append(row)
    return headers, rows


def main():
    md_text = INPUT_PATH.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1

        elif re.match(r"^表\d+-\d+\s+", line):
            add_table_caption(doc, line.strip())
            i += 1

        elif line.startswith("|") and i + 2 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            add_table(doc, headers, rows)

        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1

        elif line.strip() == "":
            i += 1

        else:
            add_paragraph(doc, line.strip())
            i += 1

    doc.save(OUTPUT_PATH)
    print(f"已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
