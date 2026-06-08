# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "docs" / "graduation"
OUTPUT_PATH = OUTPUT_DIR / "DocFlow_毕业设计任务书_精简版.docx"

TITLE = "毕业设计（论文）任务书"
TOPIC = "基于Python的多格式文档自动化处理与内容提取工具设计"


def ensure_rfonts(run, east: str = "宋体", latin: str = "Times New Roman") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)


def set_run_font(run, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    ensure_rfonts(run)


def set_paragraph_format(paragraph, *, first_indent: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(24) if first_indent else Pt(0)


def add_para(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    align=None,
    first_indent: bool = True,
    size: float = 12,
):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    set_paragraph_format(paragraph, first_indent=first_indent)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    add_para(doc, text, bold=True, first_indent=False)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def format_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph, first_indent=False)
                for run in paragraph.runs:
                    set_run_font(run)


def append_items(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item)


def build_document() -> Document:
    doc = Document()
    set_doc_defaults(doc)

    add_para(
        doc,
        TITLE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    add_para(doc, "", first_indent=False)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    info = [
        ("课题名称", TOPIC),
        ("课题类型", "工程实践类 / 软件工程系统设计与实现"),
        ("学生姓名", "待填写"),
        ("指导教师", "待填写"),
    ]
    for row, (key, value) in zip(table.rows, info):
        row.cells[0].text = key
        row.cells[1].text = value
    format_table(table)
    add_para(doc, "", first_indent=False)

    add_heading(doc, "一、主要内容")
    append_items(
        doc,
        [
            "本课题以 DocFlow 项目为基础，设计并实现一套基于 Python 的多格式文档自动化处理与内容提取工具。系统面向常见办公文档处理场景，支持 PDF、Word、Excel、PPT、TXT、Markdown、CSV、JSON 及图片等文件上传，能够完成正文、表格、元数据、关键词、摘要及 OCR 文本的提取，并将结果以 TXT、JSON、Markdown、CSV 等格式输出。",
            "学生需完成系统需求分析、总体设计、核心功能实现、测试评价和论文撰写。主要工作包括：构建 Flask 后端接口和前端交互页面；设计多格式文档解析流程和 OCR 兜底流程；完成内容提取、结果展示、批量测试、报告生成和轻量化部署说明；结合现有票据样本，将发票字段抽取作为结构化内容提取的典型验证场景，但不得将课题表述为专门的财税系统。",
            "系统应形成可运行、可演示、可复现的工程成果。论文中应重点说明文档处理流程、OCR 集成方式、内容提取策略、模块划分、接口设计、测试结果、局限性与改进方向，不夸大为自研 OCR 模型或商业级文档智能平台。",
        ],
    )

    add_heading(doc, "二、基本要求")
    append_items(
        doc,
        [
            "1. 覆盖毕业要求 3.4：能够在社会、健康、安全、法律、文化和环境等现实约束下评价并优化系统。学生应分析上传文档中可能包含的姓名、单位、票据号码、金额等敏感信息，提出临时文件清理、日志脱敏、授权样本使用和导出控制等措施，并比较 OCR 准确率、运行速度、CPU / 内存占用和部署成本。",
            "2. 覆盖毕业要求 6.3 和 8.3：能够分析软件工程实践对社会、安全、法律及文化的影响，并理解软件工程师应承担的公众安全、健康和环境保护责任。学生应说明错误识别、误导性摘要、字段误填和数据泄露可能造成的影响，明确系统输出仅作辅助参考，重要文档结论需人工复核。",
            "3. 覆盖毕业要求 7.2：能够评价软件工程实践对环境和社会可持续发展的影响。学生应说明文档数字化和自动提取对减少纸质流转、降低重复录入的价值，同时分析 OCR 计算资源、云部署能耗、缓存文件和历史报告堆积带来的成本，提出轻量部署、按需识别和定期清理策略。",
            "4. 覆盖毕业要求 10.1 和 11.3：能够通过模型、文档报告、代码等方式与用户和同行沟通，并在复杂软件项目中应用项目管理和经济决策方法。学生需提交需求说明、架构图、流程图、接口说明、测试报告、部署说明、论文和答辩 PPT，并能说明技术选型、任务分解、进度控制、风险应对和成本取舍。",
            "5. 时间节点要求：第 1-2 周完成文献翻译和文献综述；第 3 周完成开题报告并参加开题答辩；第 4-7 周完成需求分析、总体设计和核心功能联调；第 8 周进行期中检查；第 9-12 周完成 OCR、内容提取、批量测试和部署验证；第 13-14 周完成论文主体和材料整理；第 15 周完成论文评阅、查重、修改和答辩准备；第 16 周参加系、院二级答辩。",
        ],
    )

    add_heading(doc, "三、思政要求")
    append_items(
        doc,
        [
            "学生应树立精益求精的工程态度，在文档解析、OCR 识别、内容提取、异常处理和测试评价中坚持用真实样本和错误案例持续改进系统，不停留于简单功能演示。",
            "学生应强化职业伦理、知识产权和数据安全意识。处理票据、合同、报表等文档时，应使用授权或脱敏样本，不公开传播含敏感信息的材料，不伪造测试数据，不夸大系统准确率和适用范围。",
            "学生应理解软件系统对社会效率和环境保护的影响，在设计中关注纸质材料数字化、重复劳动减少、计算资源节约、临时文件清理和绿色低碳部署，体现软件工程服务社会的责任意识。",
        ],
    )

    add_heading(doc, "四、参考资料")
    references = [
        "[1] McKinney W. Python for Data Analysis[M]. 3rd ed. Sebastopol: O'Reilly Media, 2022.",
        "[2] Huang Y, Lv T, Cui L, Lu Y, Wei F. LayoutLMv3: Pre-training for Document AI with Unified Text and Image Masking[C]. Proceedings of the 30th ACM International Conference on Multimedia. New York: ACM, 2022.",
        "[3] Li C, Liu W, Guo R, Yin X, Jiang K, Du Y, Du Y, Zhu L, Lai B, Hu X, Yu D, Ma Y. PP-OCRv3: More Attempts for the Improvement of Ultra Lightweight OCR System[J/OL]. arXiv:2206.03001, 2022.",
        "[4] Blecher L, Cucurull G, Scialom T, Stojnic R. Nougat: Neural Optical Understanding for Academic Documents[C]. Proceedings of the Twelfth International Conference on Learning Representations. ICLR, 2024.",
        "[5] Pallets. Flask Documentation[EB/OL]. Pallets Projects, 2026[2026-05-24]. https://flask.palletsprojects.com/.",
    ]
    for reference in references:
        add_para(doc, reference, first_indent=False, size=10.5)

    for paragraph in doc.paragraphs:
        is_ref = paragraph.text.startswith("[")
        is_title = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        is_heading = bool(paragraph.runs and paragraph.runs[0].bold and not is_title)
        set_paragraph_format(
            paragraph,
            first_indent=not (is_ref or is_title or is_heading or paragraph.text == ""),
        )
        for run in paragraph.runs:
            set_run_font(run, size=10.5 if is_ref else 12, bold=run.bold)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
